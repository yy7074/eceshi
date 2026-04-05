#!/usr/bin/env python3
"""读取 docx 文件，提取文字和图片信息，按顺序输出"""

import os
import sys
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def extract_docx(filepath, output_dir="docx_images"):
    doc = Document(filepath)
    os.makedirs(output_dir, exist_ok=True)

    # 提取所有图片
    image_count = 0
    image_map = {}  # rId -> filename
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
            ext = os.path.splitext(rel.target_ref)[1]
            fname = f"image_{image_count}{ext}"
            image_map[rel.rId] = fname
            with open(os.path.join(output_dir, fname), "wb") as f:
                f.write(rel.target_part.blob)

    print(f"=== 文档：{os.path.basename(filepath)} ===")
    print(f"=== 共提取 {image_count} 张图片到 {output_dir}/ ===\n")

    # 按段落顺序输出文字和图片
    from docx.oxml.ns import qn
    para_num = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        # 检查段落中是否有图片
        images_in_para = []
        for run in para.runs:
            for drawing in run._element.findall(qn('w:drawing')):
                for blip in drawing.iter(qn('a:blip')):
                    rId = blip.get(qn('r:embed'))
                    if rId and rId in image_map:
                        images_in_para.append(image_map[rId])
            # 也检查 w:pict 中的图片（旧格式）
            for pict in run._element.findall(qn('w:pict')):
                for imagedata in pict.iter(qn('v:imagedata')):
                    rId = imagedata.get(qn('r:id'))
                    if rId and rId in image_map:
                        images_in_para.append(image_map[rId])

        if text or images_in_para:
            para_num += 1
            if text:
                print(f"[段落{para_num}] {text}")
            for img in images_in_para:
                print(f"[段落{para_num}] [图片: {output_dir}/{img}]")

    # 也检查表格
    for i, table in enumerate(doc.tables):
        print(f"\n--- 表格 {i+1} ---")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  行{row_idx+1}: {' | '.join(cells)}")

if __name__ == "__main__":
    filepath = "/Users/yy/Documents/GitHub/eceshi/后端所有待完善问题20260213.docx"
    extract_docx(filepath, "/Users/yy/Documents/GitHub/eceshi/docx_images")
